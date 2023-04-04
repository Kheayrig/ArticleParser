import os
import re

from bs4 import NavigableString, Tag
from docx import Document
from docx.shared import Pt, Cm

from app.ArticleParser.Formatters.FormatSettings import DocxSettings


class ToDocxFormatter:
    """ Преобразует теги в документ .docx

                Attributes:
                    document (Document): документ .docx
                    settings (DocxSettings): файл настроек
                    font_size (str): размер шрифта = 14
                    paragraph_len (Cm): ширина абзаца
        """

    def __init__(self, settings=DocxSettings()):
        self.document = Document()
        self.settings = settings
        style = self.document.styles['Normal']
        style.font.name = settings.font
        self.font_size = settings.font_size
        style.font.size = Pt(self.font_size)
        self.paragraph_len = self.convert_symbols_to_cm()

    def convert_symbols_to_cm(self):
        """
        Конвертирует ширину строки из символов в см
        :return: Cm
        """
        return Cm((self.settings.str_len * self.font_size) / 28.3465)

    def save_file(self, path: str, filename: str) -> str:
        """
        Сохраняет self.document
        :param path: путь к директории файла
        :param filename: имя файла
        :return: full_path - полный путь
        """
        full_path = os.path.join(path, f'{filename}.docx')
        self.document.save(full_path)
        return full_path

    def format_contents(self, contents: list):
        """
        Проходит по тегам, создаёт параграфы и добавляет разделитель в конец абзаца
        :param contents: теги
        :return: Document
        """
        for i in range(len(contents)):
            p = self.document.add_paragraph()
            match contents[i].name:
                case 'div' | 'span' | 'p':
                    p.width = self.paragraph_len
                    children = contents[i].contents
                    self.format_children(children, p)
                    p.add_run(self.settings.p_separator)
                case _:
                    self.format_tag(contents[i], p)
        return self.document

    def format_children(self, contents: list, p) -> None:
        """
        Проходит по дочерним элементам тега: вызывает метод класса по названию тега или печатает строку абзаца
        :param contents: список Tag | NavigableString
        :param p: текущий абзац
        :return: None
        """
        for j in range(len(contents)):
            if isinstance(contents[j], NavigableString):
                p.add_run(contents[j].text)
                p.add_run(' ')
                continue
            self.format_tag(contents[j], p)

    def format_tag(self, tag, p) -> None:
        """
        Определяет тег по названию и вызывает соответствующий метод класс, иначе добавляет текст в текущий абзац
        :param tag: тег
        :param p: текущий абзац
        :return: None
        """
        match tag.name:
            case 'ul' | 'ol':
                self.ul(tag)
            case 'h1':
                self.h(tag, 6, p)
            case 'h2':
                self.h(tag, 4, p)
            case 'h3':
                self.h(tag, 2, p)
            case 'a':
                self.a(tag, p)
            case 'b' | 'strong':
                self.b(tag, p)
            case 'i' | 'em':
                self.i(tag, p)
            case 'u' | 'ins':
                self.u(tag, p)
            case 's' | 'del':
                self.s(tag, p)
            case 'sub':
                self.sub(tag, p)
            case 'sup':
                self.sup(tag, p)
            case 'br':
                p.add_run('\n')
            case _:
                p.add_run(tag.text)
                p.add_run(' ')

    def a(self, tag: Tag, p) -> None:
        href = tag.get('href')
        if not self.settings.a_has_link or href is None:
            p.add_run(f'{tag.text} ')
        else:
            p.add_run(
                f"{tag.text} {self.settings.link_brackets.value[0]}{href}{self.settings.link_brackets.value[-1]}")

    def b(self, tag: Tag, p) -> None:
        p.add_run(tag.text).font.bold = True

    def i(self, tag: Tag, p) -> None:
        p.add_run(tag.text).font.italic = True

    def u(self, tag: Tag, p) -> None:
        p.add_run(tag.text).font.underline = True

    def s(self, tag: Tag, p) -> None:
        p.add_run(tag.text).font.strike = True

    def ul(self, tag) -> None:
        contents = tag.contents
        for li in contents:
            self.li(li)
        self.document.add_paragraph().width = self.paragraph_len

    def sub(self, tag: Tag, p) -> None:
        p.add_run(tag.text).font.subscript = True

    def sup(self, tag: Tag, p) -> None:
        p.add_run(tag.text).font.superscript = True

    def li(self, tag) -> None:
        if re.match('\s+', tag.text):
            return
        p = self.document.add_paragraph('• ')
        p.width = self.paragraph_len
        if isinstance(tag, NavigableString):
            p.add_run(tag.text)
            return
        children = tag.contents
        for item in children:
            self.format_tag(item, p)
        p.add_run('\n')

    def h(self, tag: Tag, size_plus, p) -> None:
        """
        форматирует h тег
        :param tag: тег
        :param size_plus: насколько нужно увеличить шрифт(6 - h1, 4 - h2, 2 - h3)
        :param p: текущий абзац
        :return: None
        """
        text = f'{tag.text}{self.settings.h_separator}'
        h = p.add_run(text)
        h.bold = True
        h.font.size = Pt(self.font_size + size_plus)
